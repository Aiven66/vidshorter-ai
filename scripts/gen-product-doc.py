#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_para(text, bold=False, size=None, align=None, color=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    return p

def add_table_row(table, cells_data, bold=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        if bold:
            run.bold = True
            run.font.size = Pt(10)
        else:
            run.font.size = Pt(10)
        if bg_color:
            shading = cell._element.get_or_add_tcPr()
            shading_elem = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): bg_color,
                qn('w:val'): 'clear'
            })
            shading.append(shading_elem)
    return row

# ==================== COVER ====================
doc.add_paragraph()
doc.add_paragraph()
add_para('Clipop AI', bold=True, size=36, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1a, 0x1a, 0x2e))
add_para('Product Overview Document', bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x55, 0x55, 0x55))
doc.add_paragraph()
add_para('AI-Powered Long-to-Short Video Clipping Platform', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x66, 0x66, 0x66))
doc.add_paragraph()
add_para('Version 1.0 | May 2026', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88, 0x88, 0x88))
add_para('https://clipopai.vercel.app', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x00, 0x66, 0xCC))

doc.add_page_break()

# ==================== TOC ====================
add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1. Product Overview',
    '2. Product Positioning',
    '3. Core Features',
    '4. Target Users',
    '5. Product Architecture',
    '6. Pricing Strategy',
    '7. Competitive Advantages',
    '8. Technical Stack',
    '9. SEO Keywords & Content Strategy',
    '10. Contact & Links',
]
for item in toc_items:
    add_para(item, size=12, space_after=6)

doc.add_page_break()

# ==================== 1. PRODUCT OVERVIEW ====================
add_heading_styled('1. Product Overview', level=1)

add_para('Clipop AI is an AI-powered video clipping platform that transforms long-form videos into viral short clips. Users simply paste a YouTube or Bilibili video link (or upload a local video file), and the AI automatically analyzes the content, identifies the most engaging highlight moments, and generates polished short video clips ready for social media.', space_after=12)

add_para('Product Name: Clipop AI', bold=True, size=12)
add_para('Tagline: Transform Long Videos into Viral Shorts', bold=True, size=12)
add_para('Website: https://clipopai.vercel.app', size=12)
add_para('Desktop App: Clipop Agent (macOS, Apple Silicon)', size=12)
add_para('Current Version: v0.9.5', size=12)

doc.add_paragraph()
add_para('Core Value Proposition:', bold=True, size=12)
add_bullet('Automatically detect video highlights using AI - no manual editing needed')
add_bullet('Support YouTube, Bilibili, and local video file uploads')
add_bullet('One-click generation of short clips optimized for social media platforms')
add_bullet('Desktop app for stable, high-speed local processing')

# ==================== 2. PRODUCT POSITIONING ====================
add_heading_styled('2. Product Positioning', level=1)

add_para('Clipop AI is positioned as an intelligent video repurposing tool for content creators, social media managers, and digital marketers. It bridges the gap between long-form content creation and the explosive demand for short-form video across platforms like TikTok, YouTube Shorts, Instagram Reels, and more.', space_after=12)

add_para('Positioning Statement:', bold=True, size=12)
add_para('Clipop AI is the fastest way to turn any long video into scroll-stopping short clips - powered by AI that understands what makes content engaging.', space_after=12, color=RGBColor(0x33, 0x33, 0x33))

add_para('Market Context:', bold=True, size=12)
add_bullet('Short-form video dominates social media: TikTok 1.5B+ MAU, YouTube Shorts 2B+ MAU, Instagram Reels 2B+ MAU')
add_bullet('73% of consumers prefer short-form video for discovering products/services')
add_bullet('Content creators who repurpose long content into shorts see up to 300% more reach')
add_bullet('Manual clipping is time-consuming: a 1-hour video typically takes 2-3 hours to clip manually')
add_bullet('AI-powered clipping reduces this to minutes, enabling creators to scale their output 10x')

add_para('Key Differentiators:', bold=True, size=12)
add_bullet('AI-first approach: Not just auto-cut, but intelligent highlight detection based on engagement scoring')
add_bullet('Multi-platform input: YouTube + Bilibili + local upload (most competitors only support YouTube)')
add_bullet('Desktop app option: Clipop Agent for stable, private local processing')
add_bullet('Bilingual focus: Full Chinese + English support, targeting both Western and Chinese-speaking markets')
add_bullet('Affordable pricing: Free tier available, paid plans starting at $9.9/month')

# ==================== 3. CORE FEATURES ====================
add_heading_styled('3. Core Features', level=1)

add_heading_styled('3.1 AI Highlight Detection', level=2)
add_para('The core AI engine analyzes video content through multiple dimensions:', space_after=8)
add_bullet('Frame-by-frame visual analysis to detect scene changes and key moments')
add_bullet('Audio analysis to identify high-energy segments, audience reactions, and emotional peaks')
add_bullet('Engagement scoring algorithm that ranks moments by viral potential')
add_bullet('Smart clip boundary detection for natural start/end points')

add_heading_styled('3.2 Multi-Source Video Input', level=2)
add_bullet('YouTube: Paste any YouTube video URL for instant processing')
add_bullet('Bilibili: Full support for Bilibili videos (bilibili.com / b23.tv links)')
add_bullet('Local Upload: Upload MP4, MOV, AVI files directly from your device')
add_bullet('Desktop Agent: Download and process videos locally for maximum stability')

add_heading_styled('3.3 Smart Clip Generation', level=2)
add_bullet('Automatic clip creation from detected highlights')
add_bullet('Each clip includes: title, summary, engagement score, thumbnail')
add_bullet('Preview clips before downloading')
add_bullet('Multiple format export options for different social platforms')
add_bullet('Real-time progress tracking with SSE streaming')

add_heading_styled('3.4 Desktop App - Clipop Agent', level=2)
add_para('Clipop Agent is the macOS desktop companion app that provides:', space_after=8)
add_bullet('Stable local video processing - no browser limitations')
add_bullet('Faster video downloads directly to your Mac')
add_bullet('Privacy-focused: process videos locally without uploading to cloud')
add_bullet('Seamless integration with the web platform')
add_bullet('Available for Apple Silicon (M1/M2/M3/M4) Macs')
add_bullet('System requirements: macOS 12.0 or later')

add_heading_styled('3.5 User Account & Credits System', level=2)
add_bullet('Email registration with verification code')
add_bullet('Google OAuth one-click login')
add_bullet('Credits-based usage model: each video processing costs 30 credits')
add_bullet('Daily credit reset at 00:00 UTC')
add_bullet('Processing history and dashboard')

add_heading_styled('3.6 Blog & Content Hub', level=2)
add_bullet('SEO-optimized blog articles about AI video clipping')
add_bullet('Content marketing articles targeting long-tail keywords')
add_bullet('Admin-managed content publishing system')

add_heading_styled('3.7 Internationalization (i18n)', level=2)
add_para('Full multi-language support across 32 languages:', space_after=8)
add_bullet('English, Simplified Chinese, Traditional Chinese')
add_bullet('Japanese, Korean, German, French, Italian, Spanish, Portuguese')
add_bullet('Hindi, Arabic, Bengali, Indonesian, Malay, Thai, Hebrew, Russian')
add_bullet('Urdu, Turkish, Vietnamese, Farsi, Marathi, Tamil, Polish, Telugu')
add_bullet('Nepali, Danish, Finnish, Dutch, Norwegian, Swedish')

# ==================== 4. TARGET USERS ====================
add_heading_styled('4. Target Users', level=1)

add_heading_styled('4.1 Primary User Segments', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Segment'
hdr[1].text = 'Description'
hdr[2].text = 'Pain Point'
hdr[3].text = 'Value from Clipop AI'
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

segments = [
    ['Content Creators', 'YouTubers, TikTokers, Instagram creators who produce long-form content', 'Spending hours manually clipping videos; inconsistent quality of clips', '10x faster clip creation; AI identifies best moments they might miss'],
    ['Social Media Managers', 'Agency professionals managing multiple brand accounts', 'Need to repurpose webinars, podcasts, interviews into social content at scale', 'Batch processing; consistent quality; multi-platform export'],
    ['Digital Marketers', 'Performance marketers using video ads and organic content', 'High cost of video editing; slow turnaround from brief to published clip', 'Self-service AI clipping; minutes instead of days; lower cost per clip'],
    ['Podcasters & Interviewers', 'Audio/video podcast hosts with long-form episodes', 'Long episodes get low engagement; need highlights for promotion', 'Auto-detect best moments; create promotional clips for social media'],
    ['Educators & Trainers', 'Online course creators, webinar hosts, corporate trainers', 'Long training videos have low completion rates; need micro-learning clips', 'Break 1-hour training into 60-second key takeaway clips'],
    ['Chinese-speaking Creators', 'Bilibili creators, Douyin creators, Xiaohongshu video creators', 'Few AI tools support Bilibili; most Western tools ignore Chinese platforms', 'Native Bilibili support; Chinese UI; local payment (Alipay)'],
]

for seg in segments:
    add_table_row(table, seg)

doc.add_paragraph()

add_heading_styled('4.2 User Personas', level=2)

add_para('Persona 1: "Sarah the YouTuber"', bold=True, size=11)
add_bullet('Age: 25-35, runs a YouTube channel with 50K-500K subscribers')
add_bullet('Creates: Weekly 30-60 minute videos (vlogs, tutorials, reviews)')
add_bullet('Pain: Spends 3-4 hours per video manually creating Shorts')
add_bullet('Goal: Wants to post 5-10 Shorts per long video to maximize reach')
add_bullet('Uses Clipop AI: Paste YouTube URL, get 5-10 highlight clips in 5 minutes')

add_para('Persona 2: "Wei the Bilibili Creator"', bold=True, size=11)
add_bullet('Age: 20-30, creates gaming and tech review content on Bilibili')
add_bullet('Creates: 20-40 minute videos with multiple interesting moments')
add_bullet('Pain: No good AI clipping tool supports Bilibili; manual clipping is tedious')
add_bullet('Goal: Wants to create Douyin/Xiaohongshu clips from Bilibili content')
add_bullet('Uses Clipop AI: Paste Bilibili URL, Chinese UI, Alipay payment')

add_para('Persona 3: "Marketing Manager Mike"', bold=True, size=11)
add_bullet('Age: 30-45, manages social media for a SaaS company')
add_bullet('Creates: Repurposes webinars, product demos, customer interviews')
add_bullet('Pain: Video editor costs $50+/clip; 2-day turnaround time')
add_bullet('Goal: Needs 20+ clips per week for organic social content')
add_bullet('Uses Clipop AI: Pro plan for unlimited clips, API access for automation')

# ==================== 5. PRODUCT ARCHITECTURE ====================
add_heading_styled('5. Product Architecture', level=1)

add_heading_styled('5.1 System Architecture', level=2)
add_bullet('Web Platform: Next.js 16 (App Router) + React 19, deployed on Vercel')
add_bullet('Desktop App: Electron 33 (Clipop Agent), available for macOS Apple Silicon')
add_bullet('Database: PostgreSQL via Supabase (users, videos, credits, subscriptions)')
add_bullet('AI Engine: LLM-based highlight detection (doubao-seed models)')
add_bullet('Video Processing: coze-coding-dev-sdk for video analysis and clip generation')
add_bullet('CDN & Hosting: Vercel Edge Network for global low-latency access')

add_heading_styled('5.2 Key User Flows', level=2)

add_para('Flow 1: Web Video Processing', bold=True, size=11)
add_bullet('User visits clipopai.vercel.app')
add_bullet('Pastes YouTube/Bilibili URL or uploads a video file')
add_bullet('Clicks "Process Video" (costs 30 credits)')
add_bullet('AI analyzes video: frame extraction, highlight detection, engagement scoring')
add_bullet('System generates short clips with titles, summaries, and thumbnails')
add_bullet('User previews clips and downloads desired ones')

add_para('Flow 2: Desktop App Processing', bold=True, size=11)
add_bullet('User downloads Clipop Agent from /download page')
add_bullet('Installs DMG on macOS, logs in with web account')
add_bullet('Pastes video URL or uses local file')
add_bullet('Video is processed locally on the Mac for maximum speed and privacy')
add_bullet('Clips are saved locally and can be uploaded to social platforms')

add_para('Flow 3: Subscription & Payment', bold=True, size=11)
add_bullet('Free users get daily credits (reset at 00:00 UTC)')
add_bullet('Upgrade via Pricing page: Starter ($9.9/mo) or Pro ($19.9/mo)')
add_bullet('International payment: Creem (Visa, Mastercard, Apple Pay, Google Pay)')
add_bullet('China payment: Alipay scan-to-pay')

# ==================== 6. PRICING STRATEGY ====================
add_heading_styled('6. Pricing Strategy', level=1)

add_para('Clipop AI uses a freemium model with credits-based usage:', space_after=12)

table2 = doc.add_table(rows=1, cols=5)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for i, text in enumerate(['Plan', 'Price', 'Daily Credits', 'Key Features', 'Target User']):
    hdr2[i].text = text
    for p in hdr2[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

pricing = [
    ['Free', '$0/month\n(CN: Free)', '100 credits/day\n(~3 videos/day)', 'Basic video clipping\n720p export\nWatermark included', 'Trial users\nCasual creators'],
    ['Starter', '$9.9/month\n(CN: 49 RMB/month)', '500 credits/day\n(~16 videos/day)', 'Priority processing\n1080p export\nNo watermark\nEmail support', 'Active creators\nSmall businesses'],
    ['Pro', '$19.9/month\n(CN: 99 RMB/month)', 'Unlimited credits', 'Fastest processing\n4K export quality\nNo watermark\nAPI access\nPriority support', 'Professionals\nAgencies\nTeams'],
]

for row_data in pricing:
    add_table_row(table2, row_data)

doc.add_paragraph()
add_para('Credit System Details:', bold=True, size=11)
add_bullet('1 video processing = 30 credits')
add_bullet('Credits reset daily at 00:00 UTC (unused credits do not carry over)')
add_bullet('Free plan: 100 credits/day = approximately 3 video processes per day')
add_bullet('Starter plan: 500 credits/day = approximately 16 video processes per day')
add_bullet('Pro plan: Unlimited credits = no processing limits')

add_para('Payment Methods:', bold=True, size=11)
add_bullet('International: Creem (Visa, Mastercard, Apple Pay, Google Pay)')
add_bullet('China: Alipay scan-to-pay')
add_bullet('All payments secured with TLS 256-bit encryption')

# ==================== 7. COMPETITIVE ADVANTAGES ====================
add_heading_styled('7. Competitive Advantages', level=1)

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
for i, text in enumerate(['Feature', 'Clipop AI', 'Opus Clip', 'Vizard AI']):
    hdr3[i].text = text
    for p in hdr3[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

comparison = [
    ['YouTube Support', 'Yes', 'Yes', 'Yes'],
    ['Bilibili Support', 'Yes', 'No', 'No'],
    ['Local Video Upload', 'Yes', 'Limited', 'Yes'],
    ['Desktop App', 'Yes (macOS)', 'No', 'No'],
    ['Chinese Language UI', 'Yes (32 languages)', 'No', 'No'],
    ['Alipay Payment', 'Yes', 'No', 'No'],
    ['AI Highlight Scoring', 'Yes', 'Yes', 'Yes'],
    ['Free Tier', 'Yes (100 credits/day)', 'Limited', 'Limited'],
    ['Starting Price', '$9.9/month', '$19/month', '$20/month'],
    ['4K Export', 'Yes (Pro)', 'No', 'Yes'],
    ['API Access', 'Yes (Pro)', 'No', 'No'],
]

for row_data in comparison:
    add_table_row(table3, row_data)

doc.add_paragraph()
add_para('Key Competitive Moats:', bold=True, size=12)
add_bullet('Bilibili integration: Only AI clipping tool that natively supports Bilibili - critical for Chinese market')
add_bullet('Desktop app: Clipop Agent provides stable local processing - unique among competitors')
add_bullet('Bilingual + payment localization: Full Chinese UI + Alipay = low friction for Chinese users')
add_bullet('Price advantage: Starting at $9.9/month vs. competitors at $19-20/month')
add_bullet('32-language support: Broadest international coverage in the category')

# ==================== 8. TECHNICAL STACK ====================
add_heading_styled('8. Technical Stack', level=1)

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = table4.rows[0].cells
for i, text in enumerate(['Layer', 'Technology', 'Details']):
    hdr4[i].text = text
    for p in hdr4[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

tech_stack = [
    ['Frontend Framework', 'Next.js 16 (App Router)', 'React 19, TypeScript 5, Server Components'],
    ['UI Components', 'shadcn/ui + Radix UI', 'Accessible, customizable component library'],
    ['Styling', 'Tailwind CSS 4', 'Utility-first CSS framework'],
    ['Database', 'PostgreSQL (Supabase)', 'Managed database with real-time capabilities'],
    ['AI Engine', 'doubao-seed models', 'LLM-based highlight detection and scoring'],
    ['Video Processing', 'coze-coding-dev-sdk', 'Video analysis and clip generation'],
    ['Desktop App', 'Electron 33 + electron-builder', 'macOS Apple Silicon native app'],
    ['Deployment', 'Vercel Edge Network', 'Global CDN, serverless functions, auto-deploy'],
    ['Authentication', 'Supabase Auth + Google OAuth', 'Email verification + social login'],
    ['Payment', 'Creem + Alipay', 'International cards + Chinese payment'],
    ['Analytics', 'PostHog', 'Product analytics and user behavior tracking'],
    ['i18n', 'Custom (32 languages)', 'Full UI translation for global markets'],
]

for row_data in tech_stack:
    add_table_row(table4, row_data)

# ==================== 9. SEO KEYWORDS ====================
add_heading_styled('9. SEO Keywords & Content Strategy', level=1)

add_heading_styled('9.1 Primary Keywords', level=2)
keywords_primary = [
    'AI video clipper',
    'AI video to shorts',
    'YouTube to shorts converter',
    'AI highlight detector',
    'video clipping tool',
    'long video to short video AI',
    'auto video clip generator',
    'AI shorts maker',
]
for kw in keywords_primary:
    add_bullet(kw)

add_heading_styled('9.2 Long-tail Keywords', level=2)
keywords_long = [
    'how to turn long YouTube videos into viral short clips with AI',
    'best AI tool to clip videos for TikTok and Instagram Reels',
    'automatically extract highlights from YouTube video',
    'AI video clipping tool that supports Bilibili',
    'free AI shorts generator from long videos',
    'convert webinar to short clips with AI',
    'AI podcast clipper for social media',
    'desktop app for AI video clipping macOS',
]
for kw in keywords_long:
    add_bullet(kw)

add_heading_styled('9.3 Chinese Keywords', level=2)
keywords_cn = [
    'AI视频剪辑工具',
    '长视频转短视频AI',
    'YouTube转Shorts工具',
    'B站视频剪辑AI',
    'AI自动提取视频亮点',
    '短视频自动生成器',
    'AI视频切片工具',
    '播客视频转短视频',
]
for kw in keywords_cn:
    add_bullet(kw)

add_heading_styled('9.4 Content Strategy Recommendations', level=2)
add_bullet('Blog articles targeting "how to" long-tail keywords (e.g., "How to Turn Long YouTube Videos into Viral Short Clips with AI")')
add_bullet('Comparison pages: "Clipop AI vs Opus Clip vs Vizard AI"')
add_bullet('Use case pages: "AI Video Clipping for Podcasters", "AI Shorts for Educators"')
add_bullet('Tutorial content: "Step-by-step guide to creating YouTube Shorts from long videos"')
add_bullet('Platform-specific landing pages: "/youtube-shorts-maker", "/tiktok-clip-maker", "/bilibili-clip-tool"')
add_bullet('Chinese market content: Bilibili-specific tutorials, Douyin optimization guides')
add_bullet('SEO-optimized blog with AI-generated articles targeting long-tail queries')

# ==================== 10. CONTACT ====================
add_heading_styled('10. Contact & Links', level=1)

add_para('Website: https://clipopai.vercel.app', size=12)
add_para('Download Page: https://clipopai.vercel.app/download', size=12)
add_para('Pricing: https://clipopai.vercel.app/pricing', size=12)
add_para('Blog: https://clipopai.vercel.app/blog', size=12)
add_para('GitHub Repository: https://github.com/Aiven66/vidshorter-ai', size=12)
add_para('Desktop App Release: https://github.com/Aiven66/vidshorter-ai/releases/tag/v0.9.5', size=12)

doc.add_paragraph()
add_para('---', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('This document is prepared for SEO strategy planning. For the latest product updates, visit the website directly.', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88, 0x88, 0x88))

output_path = os.path.expanduser('~/Desktop/Clipop_AI_Product_Overview.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
