#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

def add_para(text, bold=False, size=None, align=None, color=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p

def add_table_row(table, cells_data, bold=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if bold:
            run.bold = True
            run.font.size = Pt(10)
        else:
            run.font.size = Pt(10)
        if bg_color:
            shading = cell._element.get_or_add_tcPr()
            shading_elem = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): bg_color,
                qn('w:val'): 'clear'
            })
            shading.append(shading_elem)
    return row

# ==================== COVER ====================
doc.add_paragraph()
doc.add_paragraph()
add_para('Clipop AI', bold=True, size=36, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1a, 0x1a, 0x2e))
add_para('产品介绍文档', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x55, 0x55, 0x55))
doc.add_paragraph()
add_para('AI 驱动的长视频转短视频智能剪辑平台', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x66, 0x66, 0x66))
doc.add_paragraph()
add_para('版本 1.0 | 2026年5月', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88, 0x88, 0x88))
add_para('https://clipopai.vercel.app', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x00, 0x66, 0xCC))

doc.add_page_break()

# ==================== TOC ====================
add_heading_styled('目录', level=1)
toc_items = [
    '1. 产品概述',
    '2. 产品定位',
    '3. 核心功能',
    '4. 目标用户群',
    '5. 产品架构',
    '6. 定价策略',
    '7. 竞争优势',
    '8. 技术栈',
    '9. SEO 关键词与内容策略',
    '10. 联系方式与链接',
]
for item in toc_items:
    add_para(item, size=12, space_after=6)

doc.add_page_break()

# ==================== 1. PRODUCT OVERVIEW ====================
add_heading_styled('1. 产品概述', level=1)

add_para('Clipop AI 是一款 AI 驱动的智能视频剪辑平台，能够将长视频自动转化为爆款短视频。用户只需粘贴 YouTube 或 B站视频链接（或上传本地视频文件），AI 即可自动分析视频内容，识别最精彩的亮点时刻，并生成可直接发布到社交媒体的短视频片段。', space_after=12)

add_para('产品名称：Clipop AI', bold=True, size=12)
add_para('产品标语：将长视频转化为爆款短视频', bold=True, size=12)
add_para('官方网站：https://clipopai.vercel.app', size=12)
add_para('桌面客户端：Clipop Agent（macOS，Apple Silicon）', size=12)
add_para('当前版本：v0.9.5', size=12)

doc.add_paragraph()
add_para('核心价值主张：', bold=True, size=12)
add_bullet('AI 自动检测视频亮点 —— 无需手动剪辑')
add_bullet('支持 YouTube、B站 及本地视频文件上传')
add_bullet('一键生成适配各社交平台的短视频片段')
add_bullet('桌面客户端提供稳定、高速的本地处理能力')

# ==================== 2. PRODUCT POSITIONING ====================
add_heading_styled('2. 产品定位', level=1)

add_para('Clipop AI 定位为面向内容创作者、社交媒体运营者和数字营销人员的智能视频二次创作工具。它弥合了长视频内容创作与短视频平台（TikTok、YouTube Shorts、Instagram Reels、抖音、小红书等）爆发式需求之间的鸿沟。', space_after=12)

add_para('定位宣言：', bold=True, size=12)
add_para('Clipop AI 是将任意长视频转化为刷屏级短视频的最快方式 —— 由理解内容吸引力的 AI 驱动。', space_after=12, color=RGBColor(0x33, 0x33, 0x33))

add_para('市场背景：', bold=True, size=12)
add_bullet('短视频主导社交媒体：TikTok 月活超15亿，YouTube Shorts 月活超20亿，Instagram Reels 月活超20亿')
add_bullet('73% 的消费者更倾向于通过短视频了解产品或服务')
add_bullet('将长视频内容二次创作为短视频的创作者，曝光量可提升高达 300%')
add_bullet('手动剪辑耗时巨大：1小时视频通常需要2-3小时手动剪辑')
add_bullet('AI 剪辑将时间缩短至几分钟，使创作者产出效率提升10倍')

add_para('核心差异化优势：', bold=True, size=12)
add_bullet('AI 优先：不仅是自动裁剪，而是基于参与度评分的智能亮点检测')
add_bullet('多平台输入：YouTube + B站 + 本地上传（大多数竞品仅支持 YouTube）')
add_bullet('桌面客户端：Clipop Agent 提供稳定、隐私的本地处理')
add_bullet('双语聚焦：完整中英文支持，同时面向西方和中文市场')
add_bullet('价格优势：提供免费版，付费版起价仅 ¥49/月')

# ==================== 3. CORE FEATURES ====================
add_heading_styled('3. 核心功能', level=1)

add_heading_styled('3.1 AI 亮点检测', level=2)
add_para('核心 AI 引擎从多个维度分析视频内容：', space_after=8)
add_bullet('逐帧视觉分析：检测场景变化和关键时刻')
add_bullet('音频分析：识别高能量片段、观众反应和情感峰值')
add_bullet('参与度评分算法：按爆款潜力对时刻进行排名')
add_bullet('智能片段边界检测：确保自然的起止点')

add_heading_styled('3.2 多来源视频输入', level=2)
add_bullet('YouTube：粘贴任意 YouTube 视频链接即可处理')
add_bullet('B站：完整支持 B站视频（bilibili.com / b23.tv 链接）')
add_bullet('本地上传：直接上传 MP4、MOV、AVI 文件')
add_bullet('桌面客户端：本地下载和处理视频，稳定性最高')

add_heading_styled('3.3 智能片段生成', level=2)
add_bullet('从检测到的亮点自动创建视频片段')
add_bullet('每个片段包含：标题、摘要、参与度评分、缩略图')
add_bullet('下载前可预览片段')
add_bullet('多种格式导出，适配不同社交平台')
add_bullet('SSE 流式实时进度追踪')

add_heading_styled('3.4 桌面客户端 - Clipop Agent', level=2)
add_para('Clipop Agent 是 macOS 桌面伴侣应用，提供：', space_after=8)
add_bullet('稳定的本地视频处理 —— 不受浏览器限制')
add_bullet('更快的视频下载速度，直接保存到 Mac')
add_bullet('隐私优先：本地处理视频，无需上传到云端')
add_bullet('与 Web 平台无缝集成')
add_bullet('支持 Apple Silicon（M1/M2/M3/M4）芯片 Mac')
add_bullet('系统要求：macOS 12.0 或更高版本')

add_heading_styled('3.5 用户账户与积分系统', level=2)
add_bullet('邮箱注册 + 验证码验证')
add_bullet('Google OAuth 一键登录')
add_bullet('积分制使用模型：每次视频处理消耗 30 积分')
add_bullet('每日 UTC 00:00 重置积分')
add_bullet('处理历史记录与用户控制台')

add_heading_styled('3.6 博客与内容中心', level=2)
add_bullet('SEO 优化的 AI 视频剪辑相关博客文章')
add_bullet('针对长尾关键词的内容营销文章')
add_bullet('管理员后台内容发布系统')

add_heading_styled('3.7 国际化（i18n）', level=2)
add_para('完整支持 32 种语言：', space_after=8)
add_bullet('英语、简体中文、繁体中文')
add_bullet('日语、韩语、德语、法语、意大利语、西班牙语、葡萄牙语')
add_bullet('印地语、阿拉伯语、孟加拉语、印尼语、马来语、泰语、希伯来语、俄语')
add_bullet('乌尔都语、土耳其语、越南语、波斯语、马拉地语、泰米尔语、波兰语、泰卢固语')
add_bullet('尼泊尔语、丹麦语、芬兰语、荷兰语、挪威语、瑞典语')

# ==================== 4. TARGET USERS ====================
add_heading_styled('4. 目标用户群', level=1)

add_heading_styled('4.1 主要用户群体', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = '用户群体'
hdr[1].text = '描述';
hdr[2].text = '痛点';
hdr[3].text = 'Clipop AI 的价值';
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

segments = [
    ['内容创作者', 'YouTuber、TikToker、Instagram 创作者，制作长视频内容', '手动剪辑耗时数小时；片段质量不一致', '剪辑速度提升10倍；AI 发现可能错过的高光时刻'],
    ['社交媒体运营', '管理多个品牌账号的代理商或企业运营人员', '需要将直播、播客、访谈批量转化为社交内容', '批量处理；质量一致；多平台导出'],
    ['数字营销人员', '使用视频广告和自然内容推广的营销人员', '视频编辑成本高；从需求到发布周期长', '自助式 AI 剪辑；分钟级交付；单条成本更低'],
    ['播客与访谈主播', '拥有长视频/音频节目的播客和访谈主持人', '长节目互动率低；需要精彩片段做推广', '自动检测精彩时刻；生成社交媒体推广片段'],
    ['教育与培训讲师', '在线课程创作者、企业培训师、研讨会主持人', '长培训视频完播率低；需要微课片段', '将1小时培训拆分为60秒知识点片段'],
    ['中文创作者', 'B站UP主、抖音创作者、小红书视频博主', '几乎没有 AI 工具支持 B站；西方工具忽略中文平台', '原生 B站支持；中文界面；支付宝支付'],
]

for seg in segments:
    add_table_row(table, seg)

doc.add_paragraph()

add_heading_styled('4.2 用户画像', level=2)

add_para('画像一："YouTuber 小莎"', bold=True, size=11)
add_bullet('年龄：25-35岁，运营一个5万-50万粉丝的 YouTube 频道')
add_bullet('创作内容：每周发布30-60分钟视频（Vlog、教程、评测）')
add_bullet('痛点：每条视频花3-4小时手动制作 Shorts')
add_bullet('目标：希望每条长视频产出5-10条 Shorts 以最大化曝光')
add_bullet('使用 Clipop AI：粘贴 YouTube 链接，5分钟获得5-10条亮点片段')

add_para('画像二："B站UP主小魏"', bold=True, size=11)
add_bullet('年龄：20-30岁，在B站创作游戏和科技评测内容')
add_bullet('创作内容：20-40分钟视频，包含多个精彩时刻')
add_bullet('痛点：没有好的 AI 剪辑工具支持B站；手动剪辑枯燥')
add_bullet('目标：希望将B站内容转化为抖音/小红书片段')
add_bullet('使用 Clipop AI：粘贴B站链接，中文界面，支付宝支付')

add_para('画像三："营销经理老马"', bold=True, size=11)
add_bullet('年龄：30-45岁，管理一家 SaaS 公司的社交媒体')
add_bullet('创作内容：将直播回放、产品演示、客户访谈二次创作')
add_bullet('痛点：视频编辑每条成本50元以上；周转时间2天')
add_bullet('目标：每周需要20+条片段用于自然流量内容')
add_bullet('使用 Clipop AI：Pro 版无限积分，API 接入实现自动化')

# ==================== 5. PRODUCT ARCHITECTURE ====================
add_heading_styled('5. 产品架构', level=1)

add_heading_styled('5.1 系统架构', level=2)
add_bullet('Web 平台：Next.js 16（App Router）+ React 19，部署于 Vercel')
add_bullet('桌面客户端：Electron 33（Clipop Agent），支持 macOS Apple Silicon')
add_bullet('数据库：PostgreSQL（Supabase 托管），管理用户、视频、积分、订阅')
add_bullet('AI 引擎：基于 LLM 的亮点检测（doubao-seed 模型）')
add_bullet('视频处理：coze-coding-dev-sdk 进行视频分析和片段生成')
add_bullet('CDN 与托管：Vercel Edge Network 全球低延迟访问')

add_heading_styled('5.2 核心用户流程', level=2)

add_para('流程一：Web 端视频处理', bold=True, size=11)
add_bullet('用户访问 clipopai.vercel.app')
add_bullet('粘贴 YouTube/B站链接或上传视频文件')
add_bullet('点击"处理视频"（消耗30积分）')
add_bullet('AI 分析视频：帧提取、亮点检测、参与度评分')
add_bullet('系统生成带标题、摘要和缩略图的短视频片段')
add_bullet('用户预览并下载所需片段')

add_para('流程二：桌面客户端处理', bold=True, size=11)
add_bullet('用户从 /download 页面下载 Clipop Agent')
add_bullet('在 macOS 上安装 DMG，使用 Web 账户登录')
add_bullet('粘贴视频链接或使用本地文件')
add_bullet('视频在 Mac 本地处理，速度最快、隐私最安全')
add_bullet('片段保存到本地，可直接上传到社交平台')

add_para('流程三：订阅与支付', bold=True, size=11)
add_bullet('免费用户每日获得积分（UTC 00:00 重置）')
add_bullet('通过定价页升级：入门版（¥49/月）或专业版（¥99/月）')
add_bullet('海外支付：Creem（Visa、Mastercard、Apple Pay、Google Pay）')
add_bullet('国内支付：支付宝扫码')

# ==================== 6. PRICING STRATEGY ====================
add_heading_styled('6. 定价策略', level=1)

add_para('Clipop AI 采用免费增值模式，基于积分的使用量计费：', space_after=12)

table2 = doc.add_table(rows=1, cols=5)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for i, text in enumerate(['方案', '价格', '每日积分', '核心功能', '目标用户']):
    hdr2[i].text = text
    for p in hdr2[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

pricing = [
    ['免费版', '¥0/月\n($0/月)', '100 积分/天\n（约3条视频/天）', '基础视频剪辑\n720p 导出\n含水印', '试用用户\n轻度创作者'],
    ['入门版', '¥49/月\n($9.9/月)', '500 积分/天\n（约16条视频/天）', '优先处理\n1080p 导出\n无水印\n邮件支持', '活跃创作者\n小型企业'],
    ['专业版', '¥99/月\n($19.9/月)', '无限积分', '最快处理速度\n4K 导出质量\n无水印\nAPI 接入\n优先支持', '专业人士\n代理商\n团队'],
]

for row_data in pricing:
    add_table_row(table2, row_data)

doc.add_paragraph()
add_para('积分体系详情：', bold=True, size=11)
add_bullet('1次视频处理 = 30 积分')
add_bullet('积分每日 UTC 00:00 重置（未用积分不结转）')
add_bullet('免费版：100 积分/天 ≈ 每天约3次视频处理')
add_bullet('入门版：500 积分/天 ≈ 每天约16次视频处理')
add_bullet('专业版：无限积分 = 无处理限制')

add_para('支付方式：', bold=True, size=11)
add_bullet('海外：Creem（Visa、Mastercard、Apple Pay、Google Pay）')
add_bullet('国内：支付宝扫码支付')
add_bullet('所有支付均采用 TLS 256位加密保护')

# ==================== 7. COMPETITIVE ADVANTAGES ====================
add_heading_styled('7. 竞争优势', level=1)

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
for i, text in enumerate(['功能', 'Clipop AI', 'Opus Clip', 'Vizard AI']):
    hdr3[i].text = text
    for p in hdr3[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

comparison = [
    ['YouTube 支持', '✅ 支持', '✅ 支持', '✅ 支持'],
    ['B站支持', '✅ 支持', '❌ 不支持', '❌ 不支持'],
    ['本地视频上传', '✅ 支持', '⚠️ 有限', '✅ 支持'],
    ['桌面客户端', '✅ macOS 客户端', '❌ 无', '❌ 无'],
    ['中文界面', '✅ 支持（32种语言）', '❌ 不支持', '❌ 不支持'],
    ['支付宝支付', '✅ 支持', '❌ 不支持', '❌ 不支持'],
    ['AI 亮点评分', '✅ 支持', '✅ 支持', '✅ 支持'],
    ['免费版', '✅ 100积分/天', '⚠️ 有限', '⚠️ 有限'],
    ['起步价', '$9.9/月（¥49/月）', '$19/月', '$20/月'],
    ['4K 导出', '✅ 专业版', '❌ 不支持', '✅ 支持'],
    ['API 接入', '✅ 专业版', '❌ 不支持', '❌ 不支持'],
]

for row_data in comparison:
    add_table_row(table3, row_data)

doc.add_paragraph()
add_para('核心竞争壁垒：', bold=True, size=12)
add_bullet('B站集成：唯一原生支持B站的 AI 剪辑工具 —— 对中国市场至关重要')
add_bullet('桌面客户端：Clipop Agent 提供稳定的本地处理 —— 竞品中独一无二')
add_bullet('双语 + 支付本地化：完整中文界面 + 支付宝 = 中国用户零门槛')
add_bullet('价格优势：¥49/月起 vs 竞品 ¥130+/月起')
add_bullet('32种语言支持：同类产品中国际化覆盖最广')

# ==================== 8. TECHNICAL STACK ====================
add_heading_styled('8. 技术栈', level=1)

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = table4.rows[0].cells
for i, text in enumerate(['层级', '技术', '说明']):
    hdr4[i].text = text
    for p in hdr4[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

tech_stack = [
    ['前端框架', 'Next.js 16（App Router）', 'React 19、TypeScript 5、服务端组件'],
    ['UI 组件', 'shadcn/ui + Radix UI', '无障碍、可定制的组件库'],
    ['样式方案', 'Tailwind CSS 4', '原子化 CSS 框架'],
    ['数据库', 'PostgreSQL（Supabase）', '托管数据库，支持实时能力'],
    ['AI 引擎', 'doubao-seed 模型', '基于 LLM 的亮点检测与评分'],
    ['视频处理', 'coze-coding-dev-sdk', '视频分析与片段生成'],
    ['桌面客户端', 'Electron 33 + electron-builder', 'macOS Apple Silicon 原生应用'],
    ['部署', 'Vercel Edge Network', '全球 CDN、Serverless 函数、自动部署'],
    ['认证', 'Supabase Auth + Google OAuth', '邮箱验证 + 社交登录'],
    ['支付', 'Creem + 支付宝', '海外信用卡 + 国内支付宝'],
    ['数据分析', 'PostHog', '产品分析与用户行为追踪'],
    ['国际化', '自研方案（32种语言）', '完整 UI 翻译覆盖全球市场'],
]

for row_data in tech_stack:
    add_table_row(table4, row_data)

# ==================== 9. SEO KEYWORDS ====================
add_heading_styled('9. SEO 关键词与内容策略', level=1)

add_heading_styled('9.1 核心关键词（英文）', level=2)
keywords_primary_en = [
    'AI video clipper - AI视频剪辑器',
    'AI video to shorts - AI视频转短视频',
    'YouTube to shorts converter - YouTube转Shorts工具',
    'AI highlight detector - AI亮点检测器',
    'video clipping tool - 视频剪辑工具',
    'long video to short video AI - 长视频转短视频AI',
    'auto video clip generator - 自动视频片段生成器',
    'AI shorts maker - AI短视频制作器',
    'AI video repurposing tool - AI视频二次创作工具',
    'shorts generator from long videos - 长视频生成短视频',
]
for kw in keywords_primary_en:
    add_bullet(kw)

add_heading_styled('9.2 核心关键词（中文）', level=2)
keywords_primary_cn = [
    'AI视频剪辑工具',
    'AI短视频制作',
    '长视频转短视频',
    'YouTube转Shorts',
    'B站视频剪辑',
    'AI视频切片',
    '视频亮点提取',
    '短视频自动生成',
    'AI视频二次创作',
    '视频剪辑AI工具',
]
for kw in keywords_primary_cn:
    add_bullet(kw)

add_heading_styled('9.3 长尾关键词（英文）', level=2)
keywords_long_en = [
    'how to turn long YouTube videos into viral short clips with AI',
    'best AI tool to clip videos for TikTok and Instagram Reels',
    'automatically extract highlights from YouTube video',
    'AI video clipping tool that supports Bilibili',
    'free AI shorts generator from long videos',
    'convert webinar to short clips with AI',
    'AI podcast clipper for social media',
    'desktop app for AI video clipping macOS',
    'AI video highlight extractor free',
    'repurpose long form video into short form content AI',
]
for kw in keywords_long_en:
    add_bullet(kw)

add_heading_styled('9.4 长尾关键词（中文）', level=2)
keywords_long_cn = [
    '如何用AI将长视频剪辑成短视频',
    'AI自动提取视频精彩片段工具',
    'YouTube长视频转Shorts最佳工具',
    'B站视频AI自动剪辑软件',
    '免费AI短视频生成器',
    '播客视频转抖音短视频AI工具',
    'AI视频切片工具哪个好',
    '长视频自动剪辑成多条短视频',
    'AI视频二次创作工具推荐',
    'Mac上最好的AI视频剪辑软件',
    '直播回放自动生成短视频AI',
    '网课视频剪辑成微课片段AI',
    '访谈视频自动提取金句片段',
    'YouTube视频批量生成TikTok内容',
    'B站UP主必备AI剪辑工具',
]
for kw in keywords_long_cn:
    add_bullet(kw)

add_heading_styled('9.5 场景关键词', level=2)
keywords_scene = [
    'YouTube Shorts制作工具 / YouTube Shorts maker',
    'TikTok视频剪辑AI / TikTok video editor AI',
    'Instagram Reels生成器 / Instagram Reels generator',
    '抖音短视频制作 / Douyin short video maker',
    '小红书视频剪辑 / Xiaohongshu video editor',
    '直播切片工具 / Live stream clipping tool',
    '播客精彩片段提取 / Podcast highlight extractor',
    '网课微课制作 / Course micro-lesson creator',
    '产品演示视频剪辑 / Product demo video editor',
    '访谈金句提取 / Interview quote extractor',
]
for kw in keywords_scene:
    add_bullet(kw)

add_heading_styled('9.6 品牌与竞品关键词', level=2)
keywords_brand = [
    'Clipop AI / Clipop视频剪辑',
    'Clipop Agent桌面版',
    'Opus Clip替代品 / Opus Clip alternative',
    'Vizard AI替代品 / Vizard AI alternative',
    '2short AI替代品 / 2short AI alternative',
    'Klap替代品 / Klap alternative',
    '比Opus Clip更便宜的视频剪辑工具',
    '支持B站的AI视频剪辑工具',
    '中文AI视频剪辑软件',
    '有免费版的AI短视频工具',
]
for kw in keywords_brand:
    add_bullet(kw)

add_heading_styled('9.7 内容策略建议', level=2)
add_bullet('博客文章：针对"如何做"类长尾关键词（如"如何用AI将长视频剪辑成短视频"）')
add_bullet('对比页面："Clipop AI vs Opus Clip vs Vizard AI 对比评测"')
add_bullet('场景落地页："播客AI剪辑"、"教育视频微课制作"、"直播切片工具"')
add_bullet('教程内容："从长视频制作YouTube Shorts的完整指南"')
add_bullet('平台专属落地页：/youtube-shorts-maker、/tiktok-clip-maker、/bilibili-clip-tool')
add_bullet('中国市场内容：B站专属教程、抖音优化指南、小红书视频策略')
add_bullet('SEO 博客：AI 生成文章针对长尾查询，持续更新内容库')
add_bullet('视频教程：制作产品使用教程视频，发布到 YouTube 和B站，形成内容闭环')

# ==================== 10. CONTACT ====================
add_heading_styled('10. 联系方式与链接', level=1)

add_para('官方网站：https://clipopai.vercel.app', size=12)
add_para('下载页面：https://clipopai.vercel.app/download', size=12)
add_para('定价页面：https://clipopai.vercel.app/pricing', size=12)
add_para('博客页面：https://clipopai.vercel.app/blog', size=12)
add_para('GitHub 仓库：https://github.com/Aiven66/vidshorter-ai', size=12)
add_para('桌面客户端下载：https://github.com/Aiven66/vidshorter-ai/releases/tag/v0.9.5', size=12)

doc.add_paragraph()
add_para('---', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('本文档供 SEO 策略规划参考。如需最新产品信息，请访问官方网站。', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88, 0x88, 0x88))

output_path = os.path.expanduser('~/Desktop/Clipop_AI_产品介绍文档.docx')
doc.save(output_path)
print(f'文档已保存至：{output_path}')
